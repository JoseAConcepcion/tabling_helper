import tkinter as tk
from tkinter import filedialog, ttk, messagebox, scrolledtext
import os
from collections import defaultdict
import re
import csv
from docx import Document
import pandas as pd
from datetime import datetime

class WordTableExtractor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Extractor de Horarios Académicos - V2.0")
        self.root.geometry("1400x800")
        
        self.data = {}
        self.tables_data = []
        self.current_file_path = ""
        self.selector_frame = None  # Control para evitar creación múltiple del selector
        
        self.setup_gui()
        
    def setup_gui(self):
        # Frame principal
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Frame superior para controles
        control_frame = tk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Botón para cargar archivo
        btn_frame = tk.Frame(control_frame)
        btn_frame.pack(fill=tk.X, pady=5)
        
        tk.Label(btn_frame, text="Archivo:").pack(side=tk.LEFT, padx=(0, 5))
        
        self.file_label = tk.Label(btn_frame, text="Ningún archivo cargado", 
                                  fg="gray", width=50, anchor="w")
        self.file_label.pack(side=tk.LEFT, padx=5)
        
        btn_load = tk.Button(btn_frame, text="📂 Seleccionar archivo", 
                           command=self.load_file, width=20)
        btn_load.pack(side=tk.LEFT, padx=5)
        
        # Frame para botones de acción
        action_frame = tk.Frame(control_frame)
        action_frame.pack(fill=tk.X, pady=5)
        
        btn_extract = tk.Button(action_frame, text="🔍 Extraer datos", 
                               command=self.extract_data, width=15)
        btn_extract.pack(side=tk.LEFT, padx=2)
        
        btn_export = tk.Button(action_frame, text="💾 Exportar Excel", 
                              command=self.export_excel, width=15)
        btn_export.pack(side=tk.LEFT, padx=2)
        
        btn_export_txt = tk.Button(action_frame, text="📄 Exportar TXT", 
                                  command=self.export_txt, width=15)
        btn_export_txt.pack(side=tk.LEFT, padx=2)
        
        btn_clear = tk.Button(action_frame, text="🗑️ Limpiar", 
                             command=self.clear_data, width=15)
        btn_clear.pack(side=tk.LEFT, padx=2)
        
        # Etiqueta de estado
        self.status_label = tk.Label(control_frame, text="🟢 Listo", fg="green", 
                                    font=("Arial", 10))
        self.status_label.pack(side=tk.RIGHT, padx=10)
        
        # Notebook para pestañas
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Pestaña 1: Vista de datos extraídos
        self.table_frame = tk.Frame(self.notebook)
        self.notebook.add(self.table_frame, text="📊 Datos Extraídos")
        
        # Treeview con scrollbars
        tree_frame = tk.Frame(self.table_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.tree = ttk.Treeview(tree_frame, show="headings", height=20)
        
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Pestaña 2: Resumen
        self.summary_frame = tk.Frame(self.notebook)
        self.notebook.add(self.summary_frame, text="📋 Resumen")
        
        self.summary_text = scrolledtext.ScrolledText(self.summary_frame, 
                                                     wrap=tk.WORD, 
                                                     font=("Courier", 10))
        self.summary_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Pestaña 3: Vista del archivo
        self.file_view_frame = tk.Frame(self.notebook)
        self.notebook.add(self.file_view_frame, text="📄 Vista del Archivo")
        
        self.file_text = scrolledtext.ScrolledText(self.file_view_frame, 
                                                  wrap=tk.WORD, 
                                                  font=("Courier", 9))
        self.file_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Pestaña 4: Estadísticas
        self.stats_frame = tk.Frame(self.notebook)
        self.notebook.add(self.stats_frame, text="📈 Estadísticas")
        
        self.stats_text = scrolledtext.ScrolledText(self.stats_frame, 
                                                   wrap=tk.WORD, 
                                                   font=("Courier", 10))
        self.stats_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
    def load_file(self):
        """Carga un archivo para procesar"""
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Word documents", "*.docx"),
                ("Todos los archivos", "*.*"),
            ],
            title="Seleccionar archivo"
        )
        
        if file_path:
            self.current_file_path = file_path
            self.file_label.config(text=os.path.basename(file_path))
            self.status_label.config(text="🟢 Archivo cargado", fg="green")
            
            # Mostrar contenido del archivo
            self.display_file_content(file_path)
    
    def display_file_content(self, file_path):
        """Muestra el contenido del archivo en la pestaña correspondiente"""
        try:
            # Leer contenido del documento Word
            doc = Document(file_path)
            content_lines = []
            
            # Leer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    content_lines.append(para.text)
            
            # Leer tablas
            for table in doc.tables:
                content_lines.append("\n" + "="*50 + " TABLA " + "="*50)
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    content_lines.append(row_text)
            
            content = '\n'.join(content_lines)
            
            self.file_text.delete(1.0, tk.END)
            self.file_text.insert(1.0, content)
            
        except Exception as e:
            self.file_text.delete(1.0, tk.END)
            self.file_text.insert(1.0, f"Error al leer el archivo: {str(e)}")
    
    def limpiar_texto(self, texto):
        """Limpia el texto de espacios extra y caracteres especiales"""
        if not texto:
            return ""
        # Reemplazar múltiples espacios por uno solo
        texto = re.sub(r'\s+', ' ', texto.strip())
        # Eliminar caracteres de control
        texto = re.sub(r'[\x00-\x1F\x7F]', '', texto)
        return texto
    
    def formatear_tabla_mejorada(self, tabla):
        """Fusiona columnas duplicadas en una tabla"""
        if not tabla:
            return tabla
        
        # Obtener encabezados de la primera fila
        encabezados = tabla[0]
        
        # Crear un diccionario para mapear encabezados a índices de columna
        encabezado_a_indices = defaultdict(list)
        
        for idx, encabezado in enumerate(encabezados):
            encabezado_limpio = self.limpiar_texto(encabezado)
            if encabezado_limpio:
                encabezado_a_indices[encabezado_limpio].append(idx)
            else:
                encabezado_a_indices[f"_vacio_{idx}"].append(idx)
        
        # Identificar columnas que se fusionarán
        columnas_a_mantener = []
        grupos_a_fusionar = {}
        
        for encabezado, indices in encabezado_a_indices.items():
            if len(indices) > 1:
                # Mantener la primera columna, las demás se fusionarán en ella
                columna_principal = indices[0]
                columnas_a_mantener.append(columna_principal)
                grupos_a_fusionar[columna_principal] = indices[1:]
            else:
                columnas_a_mantener.append(indices[0])
        
        # Ordenar las columnas a mantener
        columnas_a_mantener.sort()
        
        # Crear nueva tabla formateada
        tabla_formateada = []
        
        for fila_idx, fila in enumerate(tabla):
            nueva_fila = []
            
            for col_idx in columnas_a_mantener:
                if col_idx >= len(fila):
                    nueva_fila.append("")
                    continue
                
                contenido_principal = self.limpiar_texto(fila[col_idx])
                
                # Verificar si esta columna tiene columnas para fusionar
                if col_idx in grupos_a_fusionar:
                    contenidos_a_fusionar = [contenido_principal]
                    
                    for dup_idx in grupos_a_fusionar[col_idx]:
                        if dup_idx < len(fila):
                            contenido_dup = self.limpiar_texto(fila[dup_idx])
                            if contenido_dup:
                                contenidos_a_fusionar.append(contenido_dup)
                    
                    # Si solo hay un contenido o todos son iguales, usar solo uno
                    if len(set(contenidos_a_fusionar)) == 1:
                        contenido_final = contenido_principal
                    else:
                        # Análisis inteligente: extraer partes únicas
                        contenido_final = self.fusionar_contenidos_inteligentemente(contenidos_a_fusionar)
                else:
                    contenido_final = contenido_principal
                
                nueva_fila.append(contenido_final)
            
            tabla_formateada.append(nueva_fila)
        
        return tabla_formateada
    
    def fusionar_contenidos_inteligentemente(self, contenidos):
        """Fusiona múltiples contenidos de forma inteligente"""
        if not contenidos:
            return ""
        
        if len(contenidos) == 1:
            return contenidos[0]
        
        if len(set(contenidos)) == 1:
            return contenidos[0]
        
        # Extraer palabras y frases comunes
        palabras_contenidos = [re.findall(r'\b\w+\b', c.lower()) for c in contenidos]
        
        # Encontrar palabras comunes a todos los contenidos
        palabras_comunes = set(palabras_contenidos[0])
        for palabras in palabras_contenidos[1:]:
            palabras_comunes = palabras_comunes.intersection(set(palabras))
        
        # Si hay muchas palabras comunes, probablemente sea el mismo contenido
        if len(palabras_comunes) > 3:
            return max(contenidos, key=len)
        
        # Buscar patrones comunes
        patrones_comunes = []
        for contenido in contenidos:
            patrones = re.findall(r'([A-Za-zÁÉÍÓÚáéíóúñÑ\.]+\s+[A-Za-zÁÉÍÓÚáéíóúñÑ\.]+)', contenido)
            patrones_comunes.append(set(patrones))
        
        # Encontrar patrones comunes a todos
        if patrones_comunes:
            patrones_interseccion = set.intersection(*patrones_comunes)
            if patrones_interseccion:
                diferencias = []
                for contenido in contenidos:
                    for patron in patrones_interseccion:
                        if patron in contenido:
                            parte_diferente = contenido.replace(patron, "").strip()
                            if parte_diferente and parte_diferente not in diferencias:
                                diferencias.append(parte_diferente)
                
                if diferencias:
                    patron_comun = next(iter(patrones_interseccion))
                    if len(diferencias) == 1:
                        return f"{patron_comun} {diferencias[0]}"
                    else:
                        return f"{patron_comun} ({' | '.join(diferencias)})"
        
        # Eliminar duplicados exactos
        contenidos_unicos = []
        for contenido in contenidos:
            if contenido and contenido not in contenidos_unicos:
                contenidos_unicos.append(contenido)
        
        if len(contenidos_unicos) == 1:
            return contenidos_unicos[0]
        else:
            return " | ".join(contenidos_unicos)
    
    def formatear_tabla_completa(self, tabla):
        """Fusiona inteligentemente columnas y filas duplicadas en una tabla"""
        if not tabla:
            return tabla
        
        # PRIMERO: Fusionar columnas duplicadas
        tabla = self.formatear_tabla_mejorada(tabla)
        
        # SEGUNDO: Identificar y fusionar filas duplicadas
        filas_fusionadas = []
        filas_por_horario = defaultdict(list)
        
        # Agrupar filas por horario (primera celda)
        for fila in tabla:
            if fila:
                horario = self.limpiar_texto(fila[0])
                filas_por_horario[horario].append(fila)
        
        # Fusionar filas con el mismo horario
        for horario, filas_grupo in filas_por_horario.items():
            if len(filas_grupo) == 1:
                filas_fusionadas.append(filas_grupo[0])
            else:
                fila_fusionada = self.fusionar_filas_duplicadas(filas_grupo)
                filas_fusionadas.append(fila_fusionada)
        
        # Ordenar las filas por horario
        filas_fusionadas = self.ordenar_filas_por_horario(filas_fusionadas)
        
        return filas_fusionadas
    
    def fusionar_filas_duplicadas(self, filas):
        """Fusiona múltiples filas con el mismo horario en una sola fila"""
        if not filas:
            return []
        
        # La primera celda (horario) es la misma para todas
        horario = filas[0][0]
        
        # Para cada columna, fusionar los contenidos de todas las filas
        num_columnas = len(filas[0])
        fila_fusionada = [horario]
        
        for col_idx in range(1, num_columnas):
            contenidos_columna = []
            
            for fila in filas:
                if col_idx < len(fila):
                    contenido = self.limpiar_texto(fila[col_idx])
                    if contenido:
                        contenidos_columna.append(contenido)
            
            if not contenidos_columna:
                contenido_final = ""
            elif len(contenidos_columna) == 1:
                contenido_final = contenidos_columna[0]
            else:
                contenido_final = self.fusionar_contenidos_inteligentemente(contenidos_columna)
            
            fila_fusionada.append(contenido_final)
        
        return fila_fusionada
    
    def ordenar_filas_por_horario(self, filas):
        """Ordena las filas por el número de horario (primera columna)"""
        if len(filas) <= 1:
            return filas
        
        def obtener_numero_horario(fila):
            if fila and fila[0]:
                match = re.search(r'^(\d+)', fila[0].strip())
                if match:
                    return int(match.group(1))
            return 0
        
        return sorted(filas, key=obtener_numero_horario)
    
    def extract_data(self):
        """Extrae y formatea las tablas del documento cargado"""
        if not self.current_file_path:
            messagebox.showwarning("Advertencia", "Por favor, carga un archivo primero.")
            return
        
        try:
            self.status_label.config(text="🟡 Extrayendo datos...", fg="orange")
            self.root.update()
            
            # Leer documento Word
            doc = Document(self.current_file_path)
            tablas_originales = []
            
            # Extraer tablas del documento
            for tabla in doc.tables:
                datos_tabla = []
                for fila in tabla.rows:
                    datos_fila = [self.limpiar_texto(celda.text) for celda in fila.cells]
                    datos_tabla.append(datos_fila)
                tablas_originales.append(datos_tabla)
            
            # Formatear cada tabla completamente
            self.tables_data = []
            for tabla in tablas_originales:
                tabla_formateada = self.formatear_tabla_completa(tabla)
                self.tables_data.append(tabla_formateada)
            
            # Mostrar datos en la interfaz
            self.display_data()
            self.update_summary()
            self.update_statistics()
            
            self.status_label.config(text="🟢 Datos extraídos correctamente", fg="green")
            messagebox.showinfo("Éxito", f"Se extrajeron {len(self.tables_data)} tabla(s) correctamente.")
            
        except Exception as e:
            self.status_label.config(text="🔴 Error en extracción", fg="red")
            messagebox.showerror("Error", f"No se pudieron extraer los datos:\n{str(e)}")
    

    def display_data(self):
        """Muestra los datos extraídos en el Treeview"""
        # Limpiar treeview
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Limpiar todas las columnas existentes
        self.tree["columns"] = ()
        
        if not self.tables_data:
            return
        
        # Mostrar la primera tabla por defecto
        tabla = self.tables_data[0]
        
        if not tabla or len(tabla) < 2:
            return
        
        # PRIMERO: Usar la primera fila como encabezados
        headers = tabla[0]
        datos = tabla[1:] if len(tabla) > 1 else []

        # Limpiar y validar encabezados
        clean_headers = []
        for i, header in enumerate(headers):
            header_str = str(header).strip()
            if not header_str:
                if i == 0:
                    clean_headers.append("Turno/Hora")
                else:
                    clean_headers.append(f"Día {i}")
            else:
                clean_headers.append(header_str)
        
        self.tree["columns"] = clean_headers
        
        for i, header in enumerate(clean_headers):
            # Las columnas de datos son #1, #2, #3, etc.
            col_id = f"#{i+1}"
            self.tree.heading(col_id, text=header, anchor="w")
            
            # Ajustar ancho según el contenido
            if i == 0:  # Primera columna (turno/hora)
                self.tree.column(col_id, width=180, minwidth=120, stretch=False, anchor="w")
            else:  # Columnas de días
                self.tree.column(col_id, width=300, minwidth=200, stretch=True, anchor="w")
        
        # Insertar SOLO los datos (no los encabezados)
        for fila_idx, fila in enumerate(datos):
            # Asegurar que la fila tenga el mismo número de columnas que los encabezados
            if len(fila) < len(clean_headers):
                # Rellenar con valores vacíos
                fila_completa = list(fila) + [""] * (len(clean_headers) - len(fila))
            elif len(fila) > len(clean_headers):
                # Recortar si tiene más columnas
                fila_completa = fila[:len(clean_headers)]
            else:
                fila_completa = list(fila)
            
            # Asegurar que todos los valores sean strings
            fila_str = []
            for val in fila_completa:
                if val is None:
                    fila_str.append("")
                else:
                    val_str = str(val).strip()
                    fila_str.append(val_str)
            
            self.tree.insert("", "end", values=fila_str)
        
        print(f"Total filas insertadas en Treeview: {len(datos)}")
        
        # Gestionar selector de tabla: destruir existente y crear nuevo si hay múltiples tablas
        if self.selector_frame is not None:
            self.selector_frame.destroy()
            self.selector_frame = None
        
        if len(self.tables_data) > 1:
            self.create_table_selector()

    def create_table_selector(self):
        """Crea un selector para cambiar entre tablas si hay múltiples (SOLO UNA VEZ)"""
        self.selector_frame = tk.Frame(self.table_frame)
        self.selector_frame.pack(fill=tk.X, padx=5, pady=5)
        
        tk.Label(self.selector_frame, text="Seleccionar tabla:").pack(side=tk.LEFT, padx=5)
        
        self.table_var = tk.StringVar()
        table_options = [f"Tabla {i+1}" for i in range(len(self.tables_data))]
        table_dropdown = ttk.Combobox(self.selector_frame, textvariable=self.table_var, 
                                     values=table_options, state="readonly", width=15)
        table_dropdown.pack(side=tk.LEFT, padx=5)
        table_dropdown.current(0)
        
        table_dropdown.bind("<<ComboboxSelected>>", self.on_table_select)
    
    def on_table_select(self, event):
        """Maneja el cambio de selección de tabla SIN crear nuevo selector"""
        if hasattr(self, 'table_var'):
            idx = int(self.table_var.get().split()[-1]) - 1
            if 0 <= idx < len(self.tables_data):
                # Limpiar treeview
                for item in self.tree.get_children():
                    self.tree.delete(item)
                
                # Limpiar todas las columnas existentes
                self.tree["columns"] = ()

                # Obtener la tabla seleccionada
                tabla = self.tables_data[idx]
                
                if not tabla or len(tabla) < 2:
                    return
                
                # Usar la primera fila como encabezados
                headers = tabla[0]
                datos = tabla[1:] if len(tabla) > 1 else []

                # Limpiar y validar encabezados
                clean_headers = []
                for i, header in enumerate(headers):
                    header_str = str(header).strip()
                    if not header_str:
                        if i == 0:
                            clean_headers.append("Turno/Hora")
                        else:
                            clean_headers.append(f"Día {i}")
                    else:
                        clean_headers.append(header_str)

                self.tree["columns"] = clean_headers
                
                for i, header in enumerate(clean_headers):
                    # Las columnas de datos son #1, #2, #3, etc.
                    col_id = f"#{i+1}"
                    self.tree.heading(col_id, text=header, anchor="w")
                    
                    # Ajustar ancho según el contenido
                    if i == 0:  # Primera columna (turno/hora)
                        self.tree.column(col_id, width=180, minwidth=120, stretch=False, anchor="w")
                    else:  # Columnas de días
                        self.tree.column(col_id, width=300, minwidth=200, stretch=True, anchor="w")
                
                # Insertar SOLO los datos (no los encabezados)
                for fila in datos:
                    # Asegurar que la fila tenga el mismo número de columnas que los encabezados
                    if len(fila) < len(clean_headers):
                        # Rellenar con valores vacíos
                        fila_completa = list(fila) + [""] * (len(clean_headers) - len(fila))
                    elif len(fila) > len(clean_headers):
                        # Recortar si tiene más columnas
                        fila_completa = fila[:len(clean_headers)]
                    else:
                        fila_completa = list(fila)

                    # Asegurar que todos los valores sean strings
                    fila_str = [str(val).strip() if val is not None else "" for val in fila_completa]
                    
                    self.tree.insert("", "end", values=fila_str)

    def update_summary(self):
        """Actualiza el resumen de las tablas extraídas"""
        self.summary_text.delete(1.0, tk.END)
        
        if not self.tables_data:
            self.summary_text.insert(tk.END, "No hay datos extraídos.")
            return
        
        self.summary_text.insert(tk.END, "RESUMEN DE TABLAS EXTRAÍDAS\n")
        self.summary_text.insert(tk.END, "="*50 + "\n\n")
        
        for i, tabla in enumerate(self.tables_data):
            self.summary_text.insert(tk.END, f"TABLA {i+1}:\n")
            self.summary_text.insert(tk.END, f"- Filas: {len(tabla)}\n")
            if tabla:
                self.summary_text.insert(tk.END, f"- Columnas: {len(tabla[0])}\n")
            
            # Analizar encabezados
            if tabla:
                self.summary_text.insert(tk.END, f"- Encabezados: {', '.join(tabla[0])}\n")
            
            self.summary_text.insert(tk.END, "\n")
    
    def update_statistics(self):
        """Actualiza las estadísticas de las tablas"""
        self.stats_text.delete(1.0, tk.END)
        
        if not self.tables_data:
            self.stats_text.insert(tk.END, "No hay datos para analizar.")
            return
        
        self.stats_text.insert(tk.END, "ESTADÍSTICAS DE TABLAS\n")
        self.stats_text.insert(tk.END, "="*50 + "\n\n")
        
        total_filas = 0
        total_columnas = 0
        
        for i, tabla in enumerate(self.tables_data):
            filas = len(tabla) - 1 if len(tabla) > 1 else 0  # Excluir encabezados
            columnas = len(tabla[0]) if tabla else 0
            total_filas += filas
            total_columnas += columnas
            
            self.stats_text.insert(tk.END, f"Tabla {i+1}:\n")
            self.stats_text.insert(tk.END, f"  • Filas de datos: {filas}\n")
            self.stats_text.insert(tk.END, f"  • Columnas: {columnas}\n")
            self.stats_text.insert(tk.END, f"  • Celdas de datos: {filas * columnas}\n")
            
            # Analizar duplicados por horario
            horarios = defaultdict(list)
            for fila_idx, fila in enumerate(tabla[1:], start=1):
                if fila and len(fila) > 0:
                    horario = fila[0].strip()
                    if horario:
                        horarios[horario].append(fila_idx)
            
            duplicados = {h: f for h, f in horarios.items() if len(f) > 1}
            if duplicados:
                self.stats_text.insert(tk.END, f"  • Horarios duplicados: {len(duplicados)}\n")
                for horario, filas in list(duplicados.items())[:3]:  # Mostrar solo primeros 3
                    self.stats_text.insert(tk.END, f"    - {horario}: {len(filas)} filas\n")
                if len(duplicados) > 3:
                    self.stats_text.insert(tk.END, f"    ... y {len(duplicados)-3} más\n")
            
            self.stats_text.insert(tk.END, "\n")
        
        self.stats_text.insert(tk.END, "TOTALES:\n")
        self.stats_text.insert(tk.END, f"• Tablas: {len(self.tables_data)}\n")
        self.stats_text.insert(tk.END, f"• Filas totales: {total_filas}\n")
        self.stats_text.insert(tk.END, f"• Columnas promedio: {total_columnas/len(self.tables_data):.1f}\n")
        self.stats_text.insert(tk.END, f"• Celdas totales: {total_filas * (total_columnas/len(self.tables_data)):.0f}\n")
    
    def export_excel(self):
        """Exporta los datos a un archivo Excel"""
        if not self.tables_data:
            messagebox.showwarning("Advertencia", "No hay datos para exportar.")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Todos los archivos", "*.*")],
            title="Guardar como Excel"
        )
        
        if file_path:
            try:
                self.status_label.config(text="🟡 Exportando a Excel...", fg="orange")
                self.root.update()
                
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    for i, tabla in enumerate(self.tables_data):
                        if tabla:
                            # Convertir a DataFrame
                            df = pd.DataFrame(tabla[1:], columns=tabla[0])
                            # Guardar en hoja de Excel
                            sheet_name = f'Tabla_{i+1}' if len(tabla[0]) <= 31 else f'T{i+1}'
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                self.status_label.config(text="🟢 Exportado a Excel", fg="green")
                messagebox.showinfo("Éxito", f"Datos exportados a:\n{file_path}")
                
            except Exception as e:
                self.status_label.config(text="🔴 Error en exportación", fg="red")
                messagebox.showerror("Error", f"No se pudo exportar a Excel:\n{str(e)}")
    
    def export_txt(self):
        """Exporta los datos a un archivo de texto"""
        if not self.tables_data:
            messagebox.showwarning("Advertencia", "No hay datos para exportar.")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("Todos los archivos", "*.*")],
            title="Guardar como TXT"
        )
        
        if file_path:
            try:
                self.status_label.config(text="🟡 Exportando a TXT...", fg="orange")
                self.root.update()
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("EXTRACCIÓN DE HORARIOS ACADÉMICOS\n")
                    f.write("="*60 + "\n\n")
                    f.write(f"Fecha de exportación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Archivo origen: {os.path.basename(self.current_file_path)}\n")
                    f.write(f"Total de tablas: {len(self.tables_data)}\n\n")
                    
                    for i, tabla in enumerate(self.tables_data):
                        f.write(f"\n{'='*60}\n")
                        f.write(f"TABLA {i+1}\n")
                        f.write(f"{'='*60}\n\n")
                        
                        if tabla:
                            # Escribir encabezados
                            headers = tabla[0]
                            f.write(" | ".join(headers) + "\n")
                            f.write("-" * (sum(len(str(h)) for h in headers) + 3 * len(headers)) + "\n")
                            
                            # Escribir datos
                            for fila in tabla[1:]:
                                f.write(" | ".join(str(cell) for cell in fila) + "\n")
                
                self.status_label.config(text="🟢 Exportado a TXT", fg="green")
                messagebox.showinfo("Éxito", f"Datos exportados a:\n{file_path}")
                
            except Exception as e:
                self.status_label.config(text="🔴 Error en exportación", fg="red")
                messagebox.showerror("Error", f"No se pudo exportar a TXT:\n{str(e)}")
    
    def clear_data(self):
        """Limpia todos los datos y restablece la interfaz"""
        self.tables_data = []
        self.current_file_path = ""
        self.file_label.config(text="Ningún archivo cargado", fg="gray")
        
        # Limpiar treeview
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Limpiar áreas de texto
        self.summary_text.delete(1.0, tk.END)
        self.stats_text.delete(1.0, tk.END)
        self.file_text.delete(1.0, tk.END)
        
        # Destruir selector de tabla si existe
        if self.selector_frame is not None:
            self.selector_frame.destroy()
            self.selector_frame = None
        
        self.status_label.config(text="🟢 Listo", fg="green")
    
    def run(self):
        """Ejecuta la aplicación"""
        self.root.mainloop()

# Función principal
def main():
    app = WordTableExtractor()
    app.run()

if __name__ == "__main__":
    main()

